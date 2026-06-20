from __future__ import annotations

import re
import os
from pathlib import Path

import fitz
import pytesseract
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter, ImageOps


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".rtf", ".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}


def _configure_tesseract() -> str:
    executable = Path(os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe"))
    if executable.exists():
        pytesseract.pytesseract.tesseract_cmd = str(executable)
    tessdata = Path(os.getenv("TESSDATA_DIR", Path(os.getenv("LOCALAPPDATA", "")) / "Tesseract-OCR" / "tessdata"))
    if tessdata.exists():
        os.environ["TESSDATA_PREFIX"] = str(tessdata)
    return "--oem 1 --psm 6 -c preserve_interword_spaces=1"


def _ocr_image(image: Image.Image) -> str:
    config = _configure_tesseract()
    image = _prepare_ocr_image(image)
    try:
        return pytesseract.image_to_string(image, lang="chi_sim+eng", config=config)
    except pytesseract.TesseractNotFoundError as exc:
        raise ValueError("OCR 需要本机安装 Tesseract") from exc
    except pytesseract.TesseractError:
        return pytesseract.image_to_string(image, lang="eng", config=config)


def _prepare_ocr_image(image: Image.Image) -> Image.Image:
    image = ImageOps.exif_transpose(image).convert("L")
    image = ImageOps.autocontrast(image)
    image = ImageEnhance.Contrast(image).enhance(1.6)
    image = image.filter(ImageFilter.SHARPEN)
    return image


def _text_quality(text: str) -> float:
    compact = re.sub(r"\s+", "", text or "")
    if not compact:
        return 0
    chinese = len(re.findall(r"[\u4e00-\u9fff]", compact))
    useful = len(re.findall(r"[\u4e00-\u9fffA-Za-z0-9]", compact))
    return useful + chinese * 1.5


def _page_text(page: fitz.Page) -> str:
    blocks = page.get_text("blocks")
    text_blocks = [
        block for block in blocks
        if len(block) >= 5 and str(block[4]).strip()
    ]
    text_blocks.sort(key=lambda block: (round(block[1] / 8), block[0]))
    return "\n".join(str(block[4]).strip() for block in text_blocks)


def _ocr_page(page: fitz.Page) -> str:
    pixmap = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
    image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
    return _ocr_image(image)


def _merge_page_text(text_layer: str, ocr_text: str) -> str:
    text_layer = text_layer.strip()
    ocr_text = ocr_text.strip()
    if not text_layer:
        return ocr_text
    if not ocr_text:
        return text_layer
    text_lines = {line.strip() for line in text_layer.splitlines() if line.strip()}
    extra_lines = [line.strip() for line in ocr_text.splitlines() if line.strip() and line.strip() not in text_lines]
    if _text_quality(ocr_text) > _text_quality(text_layer) * 1.25:
        return "\n".join([ocr_text, text_layer])
    if len(extra_lines) >= 3:
        return "\n".join([text_layer, *extra_lines])
    return text_layer


def _docx_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        with fitz.open(path) as document:
            pages = []
            for page_index, page in enumerate(document, 1):
                text_layer = _page_text(page)
                needs_ocr = _text_quality(text_layer) < 120 or bool(page.get_images(full=True))
                ocr_text = _ocr_page(page) if needs_ocr else ""
                pages.append(f"--- 第 {page_index} 页 ---\n{_merge_page_text(text_layer, ocr_text)}")
            text = "\n\n".join(pages)
    elif suffix == ".docx":
        document = Document(path)
        text = _docx_text(document)
    elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}:
        try:
            with Image.open(path) as image:
                text = _ocr_image(image)
        except pytesseract.TesseractNotFoundError as exc:
            raise ValueError("图片 OCR 需要本机安装 Tesseract") from exc
    elif suffix in {".txt", ".md", ".rtf"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"不支持的文件类型：{suffix or '未知'}")
    text = clean_and_anonymize(text)
    if len(text) < 20:
        raise ValueError("未提取到足够的简历文本")
    return text


def clean_and_anonymize(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[\t ]+", " ", text)
    text = _repair_line_breaks(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"(?<!\d)1[3-9]\d{9}(?!\d)", "[手机号已脱敏]", text)
    text = re.sub(r"[\w.+-]+@[\w-]+\.[\w.-]+", "[邮箱已脱敏]", text)
    text = re.sub(r"(?<!\d)\d{17}[\dXx](?!\d)", "[身份证号已脱敏]", text)
    return text.strip()


def _repair_line_breaks(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    repaired: list[str] = []
    for line in lines:
        if not line:
            repaired.append("")
            continue
        if repaired and _should_join_lines(repaired[-1], line):
            repaired[-1] = f"{repaired[-1]}{line}"
        else:
            repaired.append(line)
    return "\n".join(repaired)


def _should_join_lines(previous: str, current: str) -> bool:
    if not previous or not current:
        return False
    if "已脱敏" in previous:
        return False
    if re.search(r"(大学|学院|研究所|博士|硕士|本科|研究方向|项目|实习|工作|荣誉|奖励)", current):
        return False
    if re.match(r"^(---|[•*@+令e]\s?|20\d{2}|[A-Za-z ]{0,20}:|[\u4e00-\u9fff]{2,8}[：:])", current):
        return False
    if re.search(r"[。；;:：.!?？)”）]$", previous):
        return False
    if len(previous) > 120 or len(current) > 120:
        return False
    return bool(re.search(r"[\u4e00-\u9fffA-Za-z,，、(（]$", previous))
